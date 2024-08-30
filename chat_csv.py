# чат с Гигачат
from langchain.schema import ChatMessage
from langchain_community.chat_models.gigachat import GigaChat
from dotenv import load_dotenv
import base64
import os
import streamlit as st
import logging
from langchain_core.tools import tool
from langchain.agents import AgentExecutor, create_gigachat_functions_agent
from streamlit_modal import Modal
from datetime import datetime, date, timezone, timedelta, time
from docx import Document
from io import BytesIO
from giga_funcs import posts_summary
import asyncio
import nest_asyncio
import pandas as pd
from fuzzywuzzy import fuzz

load_dotenv()
nest_asyncio.apply()

csv_file = 'Client_news.csv'
document_path = 'filtered_news.docx'

logging.basicConfig(format='%(asctime)s - %(name)s - %(levelname)s - %(message)s', level=logging.INFO)
logger = logging.getLogger(__name__)

client_id = os.environ['CLIENT_ID_CM']
secret = os.environ['CLIENT_SECRET_CM']

credentials = f"{client_id}:{secret}"
auth= base64.b64encode(credentials.encode('utf-8')).decode('utf-8')

modal = Modal("Исторический поиск", key="history_search")

@tool
def news()-> None:
    """Показывает новостной фон."""
    return None

tools=[news]

giga = GigaChat(credentials=auth,
                model= 'GigaChat',
                verify_ssl_certs=False,
                scope='GIGACHAT_API_CORP',
                profanity_check=False
                )

chat=giga.bind_tools(tools=tools)

# if "document_path" not in st.session_state:
#     st.session_state.document_path=[]

if "messages" not in st.session_state:
    st.session_state.messages = [
        ChatMessage(
            role="system",
            content="Ты ассистент, который всегда готов помочь пользователю.",
        ),
        ChatMessage(
            role="assistant",
            content="Как я могу помочь вам?",
            additional_kwargs={"render_content": "Как я могу помочь вам?"},
        ),
    ]


def clear_user_data():
    st.session_state['keyword_h'] = []
    st.session_state['start_date'] = []
    st.session_state['end_date'] = []
    # st.session_state['document_path']=[]
  


def collect_news(csv_file, keyword_h, start_date, end_date, document_path, threshold=76):
    # Загрузка CSV файла
    df = pd.read_csv(csv_file, sep=',', usecols=['date', 'text'])

    # Преобразование столбца с датами в формат datetime
    df['date'] = pd.to_datetime(df['date'], format='%m/%d/%Y %I:%M:%S %p')

    # Функция для проверки совпадения по ключевому слову с учетом порога схожести
    def is_keyword_in_news(news_text, keyword_h, threshold):
        if keyword_h=='all_posts':
            return True
        else:
            return fuzz.partial_ratio(keyword_h.lower(), news_text.lower()) >= threshold


    # Фильтрация по ключевому слову и диапазону дат  
    filtered_df = df[df['date'].between(pd.to_datetime(start_date), pd.to_datetime(end_date)) & 
                     df['text'].apply(is_keyword_in_news, keyword_h=keyword_h, threshold=threshold)]


    # Создание нового DOCX документа
    doc = Document()

    # Добавление новостей в документ
    for index, row in filtered_df.iterrows():
        doc.add_heading(row['date'].strftime('%Y-%m-%d'), level=2)
        doc.add_paragraph(row['text'])
        doc.add_paragraph('----------')  # Пустая строка для разделения новостей

    # Сохранение DOCX файла
    doc.save(document_path)


if modal.is_open():
    with modal.container():
        keyword_h = st.text_input("Введите ключевое слово")
        if st.checkbox('Поиск всех постов', key='keyword_search') == True:
            keyword_h = 'all_posts'
        
        period = st.selectbox("Период поиска", ["Один день", "Неделя", "Месяц", "Произвольная дата"], key="period_selectbox")
        
        if period=='Произвольная дата':
            start_date= st.date_input('Введите дату начала поиска', key='start_date_input')

        end_date = st.date_input('Введите дату завершения поиска', key='end_date_input')
        end_date = datetime.combine(end_date, time(23, 59, 59))
   
        
        if period == 'Один день':
             start_date = end_date - timedelta(days=1)
        elif period == 'Неделя':
            start_date = end_date - timedelta(days=7)
        elif period == 'Месяц':
             start_date = end_date - timedelta(days=30)

        if st.button("Начать поиск"):
            st.session_state['keyword_h']=keyword_h
            st.session_state['start_date']=start_date
            st.session_state['end_date']=end_date
            collect_news(csv_file, st.session_state['keyword_h'], st.session_state['start_date'], st.session_state['end_date'], document_path)
            if keyword_h == 'all_posts':
                 sum_summary=asyncio.run(posts_summary(document_path, prompt_type='prompt_all_posts', keyword=None))
            else:
                 sum_summary=asyncio.run(posts_summary(document_path, prompt_type='prompt_key_word', keyword=keyword_h))
            os.remove(document_path)
            clear_user_data()
            st.session_state.message.content += sum_summary
            st.session_state.message.additional_kwargs['render_content']= sum_summary
            modal.close()
            
            
            

# Display chat messages from history on app rerun
for message in st.session_state.messages:
    with st.chat_message(message.role):
        if message.role == "assistant":
            st.markdown(message.additional_kwargs["render_content"], True)
        else:
            st.markdown(message.content, True)

if prompt := st.chat_input():
    
    message = ChatMessage(role="user", content=prompt)
    st.session_state.messages.append(message)
    
    with st.chat_message(message.role):
        st.markdown(message.content)

    message = ChatMessage(
    role="assistant", content="", additional_kwargs={"render_content": ""}
    )
    st.session_state.messages.append(message)
    

    with st.chat_message(message.role):
        message_placeholder = st.empty()
        for chunk in chat.stream(st.session_state.messages):
            if additional_kwargs := chunk.additional_kwargs:
                function_call=additional_kwargs.get('function_call')
                if function_call.get('name')=='news':
                    if ['message'] not in st.session_state:
                        st.session_state['message']=message
                    print(function_call.get('name'))
                    modal.open()   

            if chunk.additional_kwargs.get("image_uuid"):
                image_uuid = chunk.additional_kwargs.get("image_uuid")
                message.additional_kwargs[
                        "render_content"
                    ] += f"""<img src="data:png;base64,{chat.get_file(image_uuid).content}" style="width: 450px; display: block; border-radius: 10px;" >"""
            else:
                
                message.additional_kwargs["render_content"] += chunk.content
            message.content += chunk.content
            message.additional_kwargs = {
                **message.additional_kwargs,
                **chunk.additional_kwargs,
            }

            message_placeholder.markdown(
                message.additional_kwargs["render_content"] + "▌", True
            )
        message_placeholder.markdown(message.additional_kwargs["render_content"], True)

    # Каждый раз, когда пользователь нажимает что-то в интерфейсе весь скрипт выполняется заново.
    # Сохраняем токен и закрываем соединения
    st.session_state.token = chat._client.token
    chat._client.close()

