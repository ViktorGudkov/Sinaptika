from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_community.document_loaders import WebBaseLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from dotenv import load_dotenv
import base64
import os
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate
from langchain.chat_models.gigachat import GigaChat
from langchain.chains import create_retrieval_chain
from docx import Document as DocxDocument
from langchain_community.embeddings.gigachat import GigaChatEmbeddings

load_dotenv()

client_id = os.environ['CLIENT_ID_CM']
secret = os.environ['CLIENT_SECRET_CM']
credentials = f"{client_id}:{secret}"
auth= base64.b64encode(credentials.encode('utf-8')).decode('utf-8')

llm = GigaChat(credentials=auth,
                model= 'GigaChat-Pro',
                verify_ssl_certs=False,
                scope='GIGACHAT_API_CORP',
                profanity_check=False
                )

prompt = ChatPromptTemplate.from_template('''Ответь на вопрос пользователя. \
Используй при этом только информацию из контекста. Если в контексте нет \
информации для ответа, сообщи об этом пользователю.
Контекст: {context}
Вопрос: {input}
Ответ:'''
)

#sources='https://minfin.gov.ru/ru/press-center/?id_4=39265-o_zameshchenii_evroobligatsii_rossiiskoi_federatsii'

#loader = WebBaseLoader(sources.split())

#docs = loader.load()

def load_docx(file_path):
    doc = DocxDocument(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

#Чтение текста из файла .docx
file_path = "infobase.docx"
docx_content = load_docx(file_path)

# Создание объекта Document для работы с LangChain
docs = Document(page_content=docx_content)


text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000,
                                              chunk_overlap=100)
split_docs = text_splitter.split_documents([docs])

model_name = "sentence-transformers/paraphrase-multilingual-mpnet-base-v2"
model_kwargs = {'device': 'cpu'}
encode_kwargs = {'normalize_embeddings': False}
# embedding = HuggingFaceEmbeddings(model_name=model_name,
#                                   model_kwargs=model_kwargs,
#                                   encode_kwargs=encode_kwargs)

embedding=GigaChatEmbeddings(
        credentials=auth,
        verify_ssl_certs=False,
        scope='GIGACHAT_API_CORP',
        profanity_check=False
    )

vector_store = FAISS.from_documents(split_docs, embedding=embedding)
vector_store.save_local('vector_store')
#vector_store=FAISS.load_local('vector_store',embedding, allow_dangerous_deserialization=True)


# Новые документы для добавления
# new_docs = [
#     Document(page_content="Текст нового документа 1"),
#     Document(page_content="Текст нового документа 2"),
#     # Добавьте столько документов, сколько нужно
# ]


# Разделяем документы на части
#new_split_docs = text_splitter.split_documents(docs)


# Добавляем новые векторные представления в хранилище
# vector_store.add_documents(new_split_docs)

# vector_store.save_local("vector_store")


embedding_retriever = vector_store.as_retriever(search_kwargs={"k": 10})

document_chain = create_stuff_documents_chain(
    llm=llm,
    prompt=prompt
    )

retrieval_chain = create_retrieval_chain(embedding_retriever, document_chain)

q1 = 'В каком случаем можно не нужна консульская легализация документов, приложенные к оферте 2?'

resp1 = retrieval_chain.invoke(
    {'input': q1}
)

print(resp1)