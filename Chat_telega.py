# чат с Гигачат
from langchain.schema import ChatMessage
from langchain_community.chat_models.gigachat import GigaChat
from dotenv import load_dotenv
import base64
import os
import asyncio
import nest_asyncio
import streamlit as st
import logging
from langchain_core.tools import tool
from langchain.agents import AgentExecutor, create_gigachat_functions_agent
from streamlit_modal import Modal
from datetime import datetime, date, timezone, timedelta, time
from zoneinfo import ZoneInfo
from docx import Document
from io import BytesIO
from telethon import TelegramClient, events
from giga_funcs import posts_summary

load_dotenv()
nest_asyncio.apply()


initial_channels = ['bcs_express', 'tinkoff_invest_official', 'markettwits', 'cbrstocks', 'finpizdec', 'headlines_for_traders', 
                    'selfinvestor', 'if_market_news', 'forbesrussia', 'bitkogan_hotline', 'thewallstreetpro', 'roflpuls', 
                    'vedomosti', 'profinansy_news','marketpowercomics','sinara_finance','pro_bonds','fm_invest','test030424']

collected_messages=[]

logging.basicConfig(format='%(asctime)s - %(name)s - %(levelname)s - %(message)s', level=logging.INFO)
logger = logging.getLogger(__name__)

client_id = os.environ['CLIENT_ID_CM']
secret = os.environ['CLIENT_SECRET_CM']
api_id = os.environ["API_ID_"]
api_hash = os.environ["API_HASH_"]
phone = '+79049584792'

credentials = f"{client_id}:{secret}"
auth= base64.b64encode(credentials.encode('utf-8')).decode('utf-8')

modal = Modal("Исторический поиск", key="history_search")

@tool
def news()-> None:
    """Показывает новостной фон телеграм."""
    return None

tools=[news]

giga = GigaChat(credentials=auth,
                model= 'GigaChat',
                verify_ssl_certs=False,
                scope='GIGACHAT_API_CORP',
                profanity_check=False
                )

chat=giga.bind_tools(tools=tools)

# agent = create_gigachat_functions_agent(chat, tools)
# agent_executor = AgentExecutor(agent=agent, tools=tools, verbose=True)

if "document_path" not in st.session_state:
    st.session_state.document_path=[]

if "messages" not in st.session_state:
    st.session_state.messages = [
        ChatMessage(
            role="system",
            content="Ты ассистент, который всегда готов помочь пользователю.",
        ),
        ChatMessage(
            role="assistant",
            content="Как я могу помочь вам?",
            additional_kwargs={"render_content": "Как я могу помочь вам?"},
        ),
    ]

if 'client' not in st.session_state:
    st.session_state['client'] = TelegramClient('anon_', api_id, api_hash)

client = st.session_state['client']

async def fetch_channel_messages(client, channel):
    first_message = None
    async for message in client.iter_messages(channel, reverse=True, offset_date=st.session_state.start_date):
        first_message = message
        break  # Прерываем после получения первого сообщения

    min_id = first_message.id if first_message else 0
    search_value = None if st.session_state.keyword_h == 'all_posts' else st.session_state.keyword_h
    async for message in client.iter_messages(channel, search=search_value, min_id=min_id):
        if st.session_state.start_date <= message.date <= st.session_state.end_date and message.text:
            collected_messages.append({
                'channel': channel,
                'date': message.date,
                'text': message.text,
                'views': message.views if message.views is not None else 'н/д'
            })
        await asyncio.sleep(0)

async def generate_document():
    document = Document()
    for msg in collected_messages:
        message_content = f"{msg['channel']} | {msg['date']}: {msg['text']} [Просмотры: {msg['views']}]"
        document.add_paragraph(message_content)
        document.add_paragraph()
    document_path = 'collected_posts_sinap.docx'
    document.save(document_path)
    return document_path


def clear_user_data():
    st.session_state['keyword_h'] = []
    st.session_state['start_date'] = []
    st.session_state['end_date'] = []
    st.session_state['document_path']=[]
    collected_messages = []

async def collect_history(keyword_h, start_date, end_date):

    collected_messages=[]
    try:
        async with TelegramClient('anon2', api_id, api_hash) as client:
            tasks = []
            for channel in initial_channels:
                print(f"Fetching messages from {channel}")
                tasks.append(fetch_channel_messages(client, channel))

            # Wait for all fetch tasks to complete
            await asyncio.gather(*tasks)
            # Generate and send document after collecting messages
            st.session_state.document_path = await generate_document()
           
            # buffer=save_docx_to_buffer(document_path
            # st.session_state.buffer=buffer

            #return st.session_state.buffer
            return st.session_state.document_path
            
            
    except Exception as e:
        print(f"Error occurred: {e}")
        st.write('Произошла ошибка. Попробуйте повторить поиск.')

if modal.is_open():
    with modal.container():
        keyword_h = st.text_input("Введите ключевое слово")
        if st.checkbox('Поиск всех постов', key='keyword_search') == True:
            keyword_h = 'all_posts'
        
        period = st.selectbox("Период поиска", ["Один день", "Неделя", "Месяц", "Произвольная дата"], key="period_selectbox")
        
        if period=='Произвольная дата':
            start_date= st.date_input('Введите дату начала поиска', key='start_date_input')

        end_date = st.date_input('Введите дату завершения поиска', key='end_date_input')
        end_date = datetime.combine(end_date, time(23, 59, 59))
        end_date = end_date.replace(tzinfo=ZoneInfo('Europe/Moscow'))
        
        if period == 'Один день':
             start_date = end_date - timedelta(days=1)
        elif period == 'Неделя':
            start_date = end_date - timedelta(days=7)
        elif period == 'Месяц':
             start_date = end_date - timedelta(days=30)

        if st.button("Начать поиск"):
            st.session_state['keyword_h']=keyword_h
            st.session_state['start_date']=start_date
            st.session_state['end_date']=end_date
            asyncio.run(collect_history(st.session_state['keyword_h'], st.session_state['start_date'], st.session_state['end_date']))
            if keyword_h == 'all_posts':
                sum_summary=asyncio.run(posts_summary(st.session_state.document_path, prompt_type='prompt_all_posts', keyword=None))
            else:
                sum_summary=asyncio.run(posts_summary(st.session_state.document_path, prompt_type='prompt_key_word', keyword=keyword_h))
            os.remove(st.session_state.document_path)
            clear_user_data()
            
            st.session_state.message.content += sum_summary
            st.session_state.message.additional_kwargs['render_content']= sum_summary
            
            modal.close()
            

# Display chat messages from history on app rerun
for message in st.session_state.messages:
    with st.chat_message(message.role):
        if message.role == "assistant":
            st.markdown(message.additional_kwargs["render_content"], True)
        else:
            st.markdown(message.content, True)

if prompt := st.chat_input():
    
    message = ChatMessage(role="user", content=prompt)
    st.session_state.messages.append(message)
    
    with st.chat_message(message.role):
        st.markdown(message.content)

    message = ChatMessage(
    role="assistant", content="", additional_kwargs={"render_content": ""}
    )
    st.session_state.messages.append(message)
    

    with st.chat_message(message.role):
        message_placeholder = st.empty()
        for chunk in chat.stream(st.session_state.messages):
            if additional_kwargs := chunk.additional_kwargs:
                function_call=additional_kwargs.get('function_call')
                if function_call.get('name')=='news':
                    if ['message'] not in st.session_state:
                        st.session_state['message']=message
                    print(function_call.get('name'))
                    modal.open()   

            if chunk.additional_kwargs.get("image_uuid"):
                image_uuid = chunk.additional_kwargs.get("image_uuid")
                message.additional_kwargs[
                        "render_content"
                    ] += f"""<img src="data:png;base64,{chat.get_file(image_uuid).content}" style="width: 450px; display: block; border-radius: 10px;" >"""
            else:
                
                message.additional_kwargs["render_content"] += chunk.content
            message.content += chunk.content
            message.additional_kwargs = {
                **message.additional_kwargs,
                **chunk.additional_kwargs,
            }

            message_placeholder.markdown(
                message.additional_kwargs["render_content"] + "▌", True
            )
        message_placeholder.markdown(message.additional_kwargs["render_content"], True)

    # Каждый раз, когда пользователь нажимает что-то в интерфейсе весь скрипт выполняется заново.
    # Сохраняем токен и закрываем соединения
    st.session_state.token = chat._client.token
    chat._client.close()

