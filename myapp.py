import streamlit as st
from telethon import TelegramClient, events
from dotenv import load_dotenv
import os
import asyncio
import nest_asyncio
import logging
from docx import Document
from telethon.tl.types import InputPhoneContact
import streamlit as st
from streamlit_modal import Modal
from telethon.tl.functions.contacts import ImportContactsRequest
from datetime import datetime, date, timezone, timedelta, time
from zoneinfo import ZoneInfo
from io import BytesIO

load_dotenv()

logging.basicConfig(format='%(asctime)s - %(name)s - %(levelname)s - %(message)s', level=logging.INFO)
logger = logging.getLogger(__name__)

api_id = os.environ["API_ID_"]
api_hash = os.environ["API_HASH_"]
phone = '+79049584792'


nest_asyncio.apply()


initial_channels = ['bcs_express', 'tinkoff_invest_official', 'markettwits', 'cbrstocks', 'finpizdec', 'headlines_for_traders', 
                    'selfinvestor', 'if_market_news', 'forbesrussia', 'bitkogan_hotline', 'thewallstreetpro', 'roflpuls', 
                    'vedomosti', 'profinansy_news','marketpowercomics','sinara_finance','pro_bonds','fm_invest','test030424']

collected_messages=[]

keyword = st.sidebar.text_input("Введите ключевое слово для мониторинга")
user_phone_number = st.sidebar.text_input("Введите номер телефона для отправки сообщений")

if 'monitoring' not in st.session_state:
    st.session_state['monitoring'] = False

if 'main_running' not in st.session_state:
    st.session_state['main_running'] = False

if 'buffer' not in st.session_state:
    st.session_state.buffer = ''

if 'client' not in st.session_state:
    st.session_state['client'] = TelegramClient('anon_', api_id, api_hash)

client = st.session_state['client']

if 'keyword_h' not in st.session_state:
    st.session_state['keyword_h'] = ''
if 'start_date' not in st.session_state:
    st.session_state['start_date'] = ''
if 'end_date' not in st.session_state:
    st.session_state['end_date'] = ''

if 'download_button_clicked' not in st.session_state:
    st.session_state.button_clicked = False

# def check_event_loop():
#     try:
#         loop = asyncio.get_running_loop()
#         if loop.is_running():
#             print("Event loop is running")
#         else:
#             print("Event loop exists but is not running")
#     except RuntimeError:
#         print("No event loop exists")

async def get_user_id_by_phone(user_phone_number):
    try:
        # Создаем объект контакта
        contact = InputPhoneContact(client_id=0, phone=user_phone_number, first_name="Temp", last_name="Contact")
        
        # Импортируем контакт
        result = await client(ImportContactsRequest([contact]))
        
        # Получаем идентификатор пользователя
        user_id = result.users[0].id
        print(f"User ID для {user_phone_number}: {user_id}")
        return user_id
    
    except Exception as e:
        print(f"Не удалось получить идентификатор пользователя: {e}")

async def get_id():

    user_id = await get_user_id_by_phone(user_phone_number)
    print(f"User ID: {user_id}")
    return user_id

async def stop_monitoring():
    await client.disconnect()
    client.remove_event_handler(start_monitoring, events.NewMessage())
    del st.session_state['client']
    st.session_state['monitoring'] = False
    st.session_state['main_running']=False

async def start_monitoring(event):
    print('testtt')
    message_text = event.message.text.lower()  # Получаем текст сообщения и приводим к нижнему регистру
    channel_entity = await event.get_chat()  # Получаем информацию о канале
    channel_name = channel_entity.username  # Название канала
    if keyword.lower() in message_text:
        message_to_show = f"Канал: {channel_name}\n{event.message.text}"  # Формируем сообщение для пересылки
        st.write(message_to_show)
      
    # Получаем ID пользователя по номеру телефона
        user_id = await get_id()
    # Отправляем сообщение
        await client.send_message(user_id, message_to_show)

# Блок функций исторического поиска

async def fetch_channel_messages(client, channel):
    first_message = None
    async for message in client.iter_messages(channel, reverse=True, offset_date=start_date):
        first_message = message
        break  # Прерываем после получения первого сообщения

    min_id = first_message.id if first_message else 0
    search_value = None if keyword_h == 'all_posts' else keyword_h
    async for message in client.iter_messages(channel, search=search_value, min_id=min_id):
        if start_date <= message.date <= end_date and message.text:
            collected_messages.append({
                'channel': channel,
                'date': message.date,
                'text': message.text,
                'views': message.views if message.views is not None else 'н/д'
            })
        await asyncio.sleep(0)

async def generate_document():
    document = Document()
    for msg in collected_messages:
        message_content = f"{msg['channel']} | {msg['date']}: {msg['text']} [Просмотры: {msg['views']}]"
        document.add_paragraph(message_content)
        document.add_paragraph()
    document_path = 'collected_posts.docx'
    document.save(document_path)
    return document_path

# async def generate_document_giga(user_id, sum_summary):
#     document = Document()
#     document.add_paragraph(sum_summary) 
#     document_path_giga = f'{user_id}_giga_summary.docx'
#     document.save(document_path_giga)
#     return document_path_giga

def clear_user_data():
    st.session_state['keyword_h'] = []
    st.session_state['start_date'] = []
    st.session_state['end_date'] = []
    collected_messages = []

async def collect_history(keyword_h, start_date, end_date):

    collected_messages=[]
    try:
        async with TelegramClient('anon2', api_id, api_hash) as client:
            tasks = []
            for channel in initial_channels:
                print(f"Fetching messages from {channel}")
                tasks.append(fetch_channel_messages(client, channel))

            # Wait for all fetch tasks to complete
            await asyncio.gather(*tasks)

             # Generate and send document after collecting messages
            document_path = await generate_document()
            
            buffer=save_docx_to_buffer(document_path)
            st.session_state.buffer=buffer
            clear_user_data()
            os.remove(document_path)
            return st.session_state.buffer
            
            
            
    except Exception as e:
        print(f"Error occurred: {e}")
        st.write('Произошла ошибка. Попробуйте повторить поиск.')

def save_docx_to_buffer(document_path):
    buffer = BytesIO()
    with open(document_path, 'rb') as f:
        buffer.write(f.read())
    buffer.seek(0)
    return buffer

async def main():
    st.session_state['main_running'] = True
    async with client:
        client.add_event_handler(start_monitoring, events.NewMessage())
        await client.start(phone=phone)
        st.session_state['monitoring'] = True
        await client.run_until_disconnected()

# Текущий текст кнопки в зависимости от состояния
button_text = "Остановить мониторинг" if st.session_state['monitoring'] else "Начать мониторинг"

if st.sidebar.button(button_text, key='telegram_monitor_button'):
    if st.session_state['monitoring']:
        asyncio.run(stop_monitoring())

        
    else:
        st.session_state['monitoring'] = True
           
    st.rerun()

modal = Modal("Исторический поиск", key="history_search")

if st.sidebar.button("Исторический поиск", key='history_search_button'):
    modal.open()
    

if modal.is_open():
    with modal.container():
        keyword_h = st.text_input("Введите ключевое слово")
        if st.checkbox('Поиск всех постов', key='keyword_search') == True:
            keyword_h = 'all_posts'
        
        period = st.selectbox("Период поиска", ["Один день", "Неделя", "Месяц", "Произвольная дата"], key="period_selectbox")
        
        if period=='Произвольная дата':
            start_date= st.date_input('Введите дату начала поиска', key='start_date_input')

        end_date = st.date_input('Введите дату завершения поиска', key='end_date_input')
        end_date = datetime.combine(end_date, time(23, 59, 59))
        end_date = end_date.replace(tzinfo=ZoneInfo('Europe/Moscow'))
        
        if period == 'Один день':
             start_date = end_date - timedelta(days=1)
        elif period == 'Неделя':
            start_date = end_date - timedelta(days=7)
        elif period == 'Месяц':
             start_date = end_date - timedelta(days=30)

        if st.button("Начать поиск"):
            st.session_state['keyword_h']=keyword_h
            st.session_state['start_date']=start_date
            st.session_state['end_date']=end_date
            asyncio.run(collect_history(st.session_state['keyword_h'], st.session_state['start_date'], st.session_state['end_date']))
            modal.close()

if st.session_state['monitoring'] and not st.session_state['main_running']:
#if st.session_state['monitoring'] == True:
    asyncio.run(main())

# Сохраняем документ в буфер
def download_clicked():
    st.session_state.button_clicked = True

if st.session_state.buffer:
    if not st.session_state.button_clicked:
        st.download_button(
            label="Download DOCX",
            data=st.session_state.buffer,
            file_name="sample.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            on_click=download_clicked
            
    )




